#!/usr/bin/env python3
"""Build the professional Option A* foundations cost summary DOCX."""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor


OUTPUT = Path("docs/foundations-option-a-star-cost-summary.docx")

NAVY = "17365D"
BLUE = "1F4E78"
TEAL = "0F6B78"
LIGHT_BLUE = "D9EAF7"
PALE_BLUE = "EEF5FA"
PALE_TEAL = "E5F2F2"
PALE_AMBER = "FFF2CC"
AMBER = "BF7B00"
LIGHT_GREY = "F2F4F7"
MID_GREY = "D7DDE5"
DARK_GREY = "40464D"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, **edges):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge_name, attrs in edges.items():
        edge = borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            borders.append(edge)
        for key, value in attrs.items():
            edge.set(qn(f"w:{key}"), str(value))


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_field(paragraph, instruction):
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), instruction)
    paragraph._p.append(field)


def set_keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def set_keep_together(paragraph):
    paragraph.paragraph_format.keep_together = True


def add_runs(paragraph, parts):
    """Add (text, bold, italic, code) tuples to a paragraph."""
    for text, bold, italic, code in parts:
        run = paragraph.add_run(text)
        run.bold = bold
        run.italic = italic
        if code:
            run.font.name = "Aptos Mono"
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor.from_string(NAVY)


def add_body(doc, parts, *, style=None, space_after=6):
    paragraph = doc.add_paragraph(style=style)
    add_runs(paragraph, parts)
    paragraph.paragraph_format.space_after = Pt(space_after)
    return paragraph


def add_callout(doc, title, body, fill=PALE_AMBER, accent=AMBER):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Cm(16.6)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=150, start=200, bottom=150, end=200)
    set_cell_border(
        cell,
        start={"val": "single", "sz": "18", "color": accent},
        top={"val": "nil"},
        bottom={"val": "nil"},
        end={"val": "nil"},
    )
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title.upper())
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(accent)
    p = cell.add_paragraph(body)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    set_keep_together(p)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_data_table(doc, headers, rows, widths, *, total_rows=(), align_right=()):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    table.rows[0].height = Cm(0.72)
    for i, (cell, label, width) in enumerate(zip(table.rows[0].cells, headers, widths)):
        cell.width = width
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(label)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(WHITE)
        run.font.size = Pt(8.5)
        if i in align_right:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_repeat_table_header(table.rows[0])

    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        is_total = row_index in total_rows
        for i, (cell, value, width) in enumerate(zip(cells, values, widths)):
            cell.width = width
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if is_total:
                set_cell_shading(cell, LIGHT_BLUE)
            elif row_index % 2:
                set_cell_shading(cell, "F8FAFC")
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            if i in align_right:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(str(value))
            run.bold = is_total
            run.font.size = Pt(8.5)
            if is_total:
                run.font.color.rgb = RGBColor.from_string(NAVY)
        prevent_row_split(table.rows[-1])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_metric_cards(doc):
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cards = [
        ("$8.20", "Compute per run", "Measured upper bound"),
        ("$691", "Year 1 monthly average", "Compute + bronze + silver"),
        ("$2,437", "Steady-state monthly", "At 4-year retention"),
    ]
    for cell, (value, label, note) in zip(table.rows[0].cells, cards):
        cell.width = Cm(5.45)
        set_cell_shading(cell, PALE_BLUE)
        set_cell_margins(cell, top=170, start=160, bottom=170, end=160)
        set_cell_border(
            cell,
            top={"val": "single", "sz": "14", "color": BLUE},
            bottom={"val": "single", "sz": "4", "color": MID_GREY},
            start={"val": "single", "sz": "4", "color": MID_GREY},
            end={"val": "single", "sz": "4", "color": MID_GREY},
        )
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(value)
        r.bold = True
        r.font.size = Pt(20)
        r.font.color.rgb = RGBColor.from_string(BLUE)
        p = cell.add_paragraph(label)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        r = p.runs[0]
        r.bold = True
        r.font.size = Pt(8.5)
        p = cell.add_paragraph(note)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.runs[0]
        r.font.size = Pt(7.5)
        r.font.color.rgb = RGBColor.from_string(DARK_GREY)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_architecture(doc):
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    steps = [
        ("1", "MWAA", "Daily schedule"),
        ("2", "Redshift", "Wait + UNLOAD"),
        ("3", "S3 bronze", "7-day handoff"),
        ("4", "AWS Glue", "Transform"),
        ("5", "S3 silver", "Cleansed output"),
    ]
    for cell, (number, title, subtitle) in zip(table.rows[0].cells, steps):
        cell.width = Cm(3.25)
        set_cell_shading(cell, PALE_TEAL)
        set_cell_margins(cell, top=120, start=100, bottom=120, end=100)
        set_cell_border(
            cell,
            top={"val": "single", "sz": "6", "color": TEAL},
            bottom={"val": "single", "sz": "6", "color": TEAL},
            start={"val": "single", "sz": "6", "color": TEAL},
            end={"val": "single", "sz": "6", "color": TEAL},
        )
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(number)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(TEAL)
        p = cell.add_paragraph(title)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(1)
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(9)
        p = cell.add_paragraph(subtitle)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        p.runs[0].font.size = Pt(7.5)
        p.runs[0].font.color.rgb = RGBColor.from_string(DARK_GREY)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor.from_string(DARK_GREY)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    title = styles["Title"]
    title.font.name = "Aptos Display"
    title.font.size = Pt(28)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(NAVY)
    title.paragraph_format.space_after = Pt(5)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Aptos"
    subtitle.font.size = Pt(12)
    subtitle.font.color.rgb = RGBColor.from_string(TEAL)
    subtitle.paragraph_format.space_after = Pt(12)

    for name, size, before, after, color in (
        ("Heading 1", 17, 16, 6, NAVY),
        ("Heading 2", 12.5, 12, 5, BLUE),
        ("Heading 3", 10.5, 10, 4, TEAL),
    ):
        style = styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    styles["List Number"].font.name = "Aptos"
    styles["List Number"].font.size = Pt(9)
    styles["List Number"].paragraph_format.space_after = Pt(4)


def configure_page(section):
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(18)
    section.bottom_margin = Mm(18)
    section.left_margin = Mm(21)
    section.right_margin = Mm(21)
    section.header_distance = Mm(8)
    section.footer_distance = Mm(8)


def add_header_footer(section):
    header = section.header
    p = header.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("FOUNDATIONS PIPELINE  |  COST ANALYSIS")
    r.bold = True
    r.font.size = Pt(7.5)
    r.font.color.rgb = RGBColor.from_string(TEAL)
    p_pr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), TEAL)
    borders.append(bottom)
    p_pr.append(borders)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("PDD-4920  |  22 June 2026  |  Page ")
    r.font.size = Pt(7.5)
    r.font.color.rgb = RGBColor.from_string(DARK_GREY)
    add_field(p, "PAGE")


def add_title_page(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(34)
    p.add_run("COST ANALYSIS").bold = True
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = RGBColor.from_string(TEAL)

    p = doc.add_paragraph("Option A*", style="Title")
    p = doc.add_paragraph("audience_activity Foundations Pipeline", style="Subtitle")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(24)
    r = p.add_run("Cost breakdown and retention decision pack")
    r.font.size = Pt(15)
    r.font.color.rgb = RGBColor.from_string(DARK_GREY)

    add_callout(
        doc,
        "Purpose",
        "Prepared for sign-off using measured consumption from the PDD-4920 sizing test. "
        "The analysis separates recurring compute, transient bronze storage, and retained silver storage.",
        fill=PALE_BLUE,
        accent=BLUE,
    )

    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    metadata = [
        ("Status", "For sign-off"),
        ("Measurement date", "19 May 2026"),
        ("Sizing volume", "4,512,730,031 rows"),
        ("Last updated", "22 June 2026"),
    ]
    for row, (label, value) in zip(table.rows, metadata):
        row.cells[0].width = Cm(4)
        row.cells[1].width = Cm(8.5)
        for cell in row.cells:
            set_cell_margins(cell, top=85, start=80, bottom=85, end=80)
            set_cell_border(cell, bottom={"val": "single", "sz": "4", "color": MID_GREY})
        row.cells[0].paragraphs[0].add_run(label).bold = True
        row.cells[0].paragraphs[0].runs[-1].font.color.rgb = RGBColor.from_string(BLUE)
        row.cells[1].paragraphs[0].add_run(value)
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(36)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Scope note")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)
    p = doc.add_paragraph(
        "Costs are EDP-specific. Glue compute applies similarly on MAP; Redshift Serverless costs are EDP-only."
    )
    p.paragraph_format.space_after = Pt(0)
    p.runs[0].font.size = Pt(8.5)
    p.runs[0].font.color.rgb = RGBColor.from_string(DARK_GREY)


def build_document():
    doc = Document()
    configure_styles(doc)
    configure_page(doc.sections[0])
    add_header_footer(doc.sections[0])
    doc.core_properties.title = "Option A* Cost Summary: audience_activity Foundations Pipeline"
    doc.core_properties.subject = "Professional cost breakdown and retention analysis"
    doc.core_properties.author = "Foundations Programme"
    doc.core_properties.keywords = "PDD-4920, AWS, Redshift, Glue, S3, cost analysis"
    doc.core_properties.comments = "Generated from foundations-option-a-star-cost-summary.md"

    add_title_page(doc)

    executive_heading = doc.add_heading("Executive summary", level=1)
    executive_heading.paragraph_format.page_break_before = True
    add_metric_cards(doc)
    add_body(
        doc,
        [("At four-year retention, silver storage represents approximately ", False, False, False),
         ("88% of the $2,437 steady-state monthly total", True, False, False),
         (". Compute and bronze storage together remain approximately $291 per month.", False, False, False)],
    )
    add_callout(
        doc,
        "Decision required before lifecycle configuration",
        "Confirm whether silver data must be retained for the four-year DMG upper bound or the 90-day operational recommendation. "
        "The 90-day option reduces steady-state silver storage from approximately $2,146 to $445 per month, a reduction of "
        "$1,701 per month (about 79% of silver cost). The corresponding all-in monthly run rate is approximately $736.",
    )

    doc.add_heading("Cost profile", level=2)
    add_data_table(
        doc,
        ["Cost view", "Monthly", "Annualised", "Basis"],
        [
            ("Year 1 average", "~$691", "~$8,292", "Silver accumulation + 30 runs/month"),
            ("4-year steady state", "~$2,437", "~$29,244", "1,460 retained silver partitions"),
            ("90-day steady state", "~$736", "~$8,832", "Operational retention recommendation"),
        ],
        [Cm(4.0), Cm(2.6), Cm(2.8), Cm(7.0)],
        total_rows=(1,),
        align_right=(1, 2),
    )
    add_body(
        doc,
        [("Annualised values are monthly run-rate extrapolations for comparison; they are not a phased cash-flow forecast.", False, True, False)],
        space_after=2,
    )

    doc.add_heading("Pipeline architecture", level=1)
    add_architecture(doc)
    add_body(
        doc,
        [("MWAA schedules the daily workflow. A SQL sensor waits for the datashare partition, Redshift Serverless unloads it to "
          "the bronze handoff, and AWS Glue transforms it into the silver-cleansed dataset.", False, False, False)],
    )

    doc.add_heading("Compute cost", level=1)
    add_data_table(
        doc,
        ["Component", "Measured consumption", "Calculation", "Per run"],
        [
            ("Redshift Serverless UNLOAD", "8,160 RPU-seconds", "8,160 / 3,600 x $0.375/RPU-hr", "$0.85"),
            ("AWS Glue transform", "16.70 DPU-hours", "16.70 x $0.44/DPU-hr", "$7.35"),
            ("Total per run", "", "", "~$8.20"),
        ],
        [Cm(4.0), Cm(4.0), Cm(5.3), Cm(2.7)],
        total_rows=(2,),
        align_right=(3,),
    )
    add_data_table(
        doc,
        ["Run frequency", "Compute cost"],
        [("Monthly (30 runs)", "~$246"), ("Annual (365 runs)", "~$2,993")],
        [Cm(11.5), Cm(4.5)],
        align_right=(1,),
    )
    add_callout(
        doc,
        "Conservative compute basis",
        "The Glue measurement came from the initial sizing run, which created the catalog table using saveAsTable. Production uses "
        "insertInto with dynamic partition overwrite, so steady-state Glue duration and cost may be modestly lower.",
        fill=PALE_BLUE,
        accent=BLUE,
    )

    doc.add_heading("Storage cost", level=1)
    doc.add_heading("Bronze: transient 7-day handoff", level=2)
    add_body(
        doc,
        [("Bronze exists only between UNLOAD and the Glue transform. A seven-day S3 lifecycle supports same-day reruns and weekend "
          "failure recovery without another UNLOAD. Intelligent-Tiering is not warranted at this retention length.", False, False, False)],
    )
    add_data_table(
        doc,
        ["Metric", "Value"],
        [
            ("Per daily partition", "~280 GB across 128 Parquet files"),
            ("Steady-state rolling window", "7 x 280 GB = 1,960 GB"),
            ("S3 Standard rate (eu-west-1)", "$0.023/GB/month"),
            ("Monthly bronze cost", "~$45/month"),
        ],
        [Cm(7.0), Cm(9.0)],
        total_rows=(3,),
        align_right=(1,),
    )
    add_body(doc, [("The rolling-window cost is reached after seven days; there is no multi-year build-up period.", False, True, False)])

    doc.add_heading("Silver: 4-year retained dataset", level=2)
    add_body(
        doc,
        [("Each daily partition writes 309 GB. S3 Intelligent-Tiering progressively moves objects as access recency changes.", False, False, False)],
    )
    add_data_table(
        doc,
        ["Access tier", "Activates", "Rate (eu-west-1)"],
        [
            ("Frequent Access", "Day 0", "$0.023/GB/month"),
            ("Infrequent Access", "Day 31", "$0.0125/GB/month"),
            ("Archive Instant Access", "Day 91", "$0.004/GB/month"),
        ],
        [Cm(6.0), Cm(3.5), Cm(6.5)],
        align_right=(2,),
    )

    doc.add_heading("Four-year steady state", level=3)
    add_data_table(
        doc,
        ["Tier", "Data / object basis", "Monthly"],
        [
            ("Frequent: last 30 days", "30 x 309 GB = 9,270 GB", "$213"),
            ("Infrequent: days 31-90", "60 x 309 GB = 18,540 GB", "$232"),
            ("Archive Instant: days 91-1,460", "1,370 x 309 GB = 423,330 GB", "$1,693"),
            ("Intelligent-Tiering monitoring", "1,460 x 2,241 files = 3.3M objects", "$8"),
            ("Total steady state", "441 TB", "~$2,146/month"),
        ],
        [Cm(5.2), Cm(7.2), Cm(3.6)],
        total_rows=(4,),
        align_right=(2,),
    )

    doc.add_heading("Silver cost build-up", level=3)
    add_data_table(
        doc,
        ["Period", "Approximate monthly silver cost"],
        [
            ("Year 1 average", "~$400/month"),
            ("End of year 1", "~$785/month"),
            ("End of year 2", "~$1,200/month"),
            ("End of year 3", "~$1,690/month"),
            ("Year 4+ steady state", "~$2,146/month"),
        ],
        [Cm(9.5), Cm(6.5)],
        total_rows=(4,),
        align_right=(1,),
    )

    doc.add_heading("Combined monthly totals", level=1)
    add_data_table(
        doc,
        ["Component", "Monthly cost"],
        [
            ("Compute: 30 daily runs", "~$246"),
            ("Bronze: 7-day steady state", "~$45"),
            ("Silver: 4-year steady state", "~$2,146"),
            ("Total at 4-year steady state", "~$2,437/month"),
            ("Total in year 1", "~$691/month average"),
        ],
        [Cm(10.5), Cm(5.5)],
        total_rows=(3, 4),
        align_right=(1,),
    )

    doc.add_heading("Assumptions, evidence and exclusions", level=1)
    notes = [
        ("Sizing baseline. ", "All figures use the PDD-4920 sizing test for dt=2026-05-19 at 4,512,730,031 rows."),
        ("Superseded estimate. ", "The RFC's earlier ~$64/month silver estimate was pre-sizing and is replaced by the measured 309 GB/day basis."),
        ("Request charges. ", "S3 PUT and GET request costs are excluded because they are a fraction of a cent per run at this scale."),
        ("Retention policy. ", "Four years is the DMG upper bound (action A6). The RFC review's operational recommendation is 90-day silver retention at approximately $445/month."),
        ("Pending platform action. ", "Bronze lifecycle configuration awaits EDP platform provisioning (action A5)."),
        ("Price basis. ", "S3 and Redshift Serverless rates are eu-west-1 list prices as at 22 June 2026 and may change."),
        ("Glue evidence. ", "The 16.70 DPU-hours measurement used 10 x G.1X workers for approximately 101 minutes and is treated as a conservative upper bound."),
        ("Redshift evidence. ", "The EDP platform team retrieved 8,160 billed RPU-seconds from CloudWatch ChargedSeconds for 9 June 2026, 21:47-21:58 UTC. It reflects dynamic scaling, not the configured 128-RPU base capacity."),
    ]
    for title, text in notes:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.keep_together = True
        r = p.add_run(title)
        r.bold = True
        p.add_run(text)

    doc.add_heading("Sign-off", level=1)
    add_callout(
        doc,
        "Approval point",
        "Confirm the silver retention period before the lifecycle policy is configured: four-year DMG upper bound or 90-day operational recommendation.",
        fill=PALE_AMBER,
        accent=AMBER,
    )
    table = doc.add_table(rows=3, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row, label in zip(table.rows, ("Decision", "Approved by", "Date")):
        row.cells[0].width = Cm(4.0)
        row.cells[1].width = Cm(12.0)
        row.cells[0].paragraphs[0].add_run(label).bold = True
        for cell in row.cells:
            set_cell_margins(cell, top=160, bottom=160)
            set_cell_border(cell, bottom={"val": "single", "sz": "6", "color": MID_GREY})
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT.resolve())


if __name__ == "__main__":
    build_document()
